from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import pytesseract
import io
import os


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


    @staticmethod
    def load(file_path: Path) -> dict:
        ext = file_path.suffix.lower()

        if ext == ".pdf":
            return DocumentLoader._load_pdf(file_path)
        elif ext == ".docx":
            return DocumentLoader._load_docx(file_path)
        elif ext == ".txt":
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError("Unsupported file type")

    MAX_OCR_PAGES = 150  # cap OCR to avoid runaway processing on large scanned PDFs

    @staticmethod
    def _ocr_page(i: int, img: Image.Image) -> tuple:
        try:
            # --oem 1 = LSTM only (faster); --psm 6 = uniform text block
            text = pytesseract.image_to_string(img, config="--oem 1 --psm 6")
        except Exception as e:
            print(f"OCR failed for page {i+1}: {e}")
            text = ""
        return i, text

    @staticmethod
    def _load_pdf(file_path: Path) -> dict:
        doc = fitz.open(file_path)
        pages = [None] * len(doc)
        to_ocr = []

        # First pass: pull embedded text and render page images for scanned
        # pages. Rendering is cheap; the actual OCR call is the slow part.
        for i, page in enumerate(doc):
            text = page.get_text()

            if not text.strip() and len(to_ocr) < DocumentLoader.MAX_OCR_PAGES:
                # 1.5x zoom — good enough for tesseract, 44% fewer pixels than 2x
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                to_ocr.append((i, img))
            else:
                pages[i] = {"page": i + 1, "text": text}

        # Second pass: OCR the scanned pages concurrently. pytesseract shells
        # out to the tesseract binary, which releases the GIL, so this
        # actually overlaps across CPU cores instead of running one page at a time.
        if to_ocr:
            max_workers = min(len(to_ocr), (os.cpu_count() or 1) * 2)
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                for i, text in executor.map(lambda args: DocumentLoader._ocr_page(*args), to_ocr):
                    pages[i] = {"page": i + 1, "text": text}

        return {
            "filename": file_path.name,
            "type": "pdf",
            "pages": pages
        }

    @staticmethod
    def _load_docx(file_path: Path) -> dict:
        """Load DOCX and split by page breaks or sections"""
        doc = Document(file_path)
        pages = []
        current_page = 1
        current_text = []
        
        for paragraph in doc.paragraphs:
            # Check if paragraph contains a page break
            if '\f' in paragraph.text or '\x0c' in paragraph.text:
                # Save current page
                if current_text:
                    pages.append({
                        "page": current_page,
                        "text": "\n".join(current_text)
                    })
                    current_page += 1
                    current_text = []
            else:
                # Add paragraph to current page
                if paragraph.text.strip():
                    current_text.append(paragraph.text)
        
        # Add the last page
        if current_text:
            pages.append({
                "page": current_page,
                "text": "\n".join(current_text)
            })
        
        # If no page breaks found, split by approximate page size
        if len(pages) == 1 and len(pages[0]["text"]) > 3000:
            pages = DocumentLoader._split_by_length(pages[0]["text"], file_path.name)
        
        # If still only one page, that's fine - it's a short document
        if not pages:
            # Fallback: treat entire document as one page
            all_text = "\n".join(p.text for p in doc.paragraphs)
            pages = [{"page": 1, "text": all_text}]

        return {
            "filename": file_path.name,
            "type": "docx",
            "pages": pages
        }
    
    @staticmethod
    def _split_by_length(text: str, filename: str, chars_per_page: int = 3000) -> list:
        """Split long text into approximate pages"""
        pages = []
        words = text.split()
        current_page = []
        current_length = 0
        page_num = 1
        
        for word in words:
            current_page.append(word)
            current_length += len(word) + 1  # +1 for space
            
            if current_length >= chars_per_page:
                pages.append({
                    "page": page_num,
                    "text": " ".join(current_page)
                })
                page_num += 1
                current_page = []
                current_length = 0
        
        # Add remaining text
        if current_page:
            pages.append({
                "page": page_num,
                "text": " ".join(current_page)
            })
        
        return pages

    @staticmethod
    def _load_txt(file_path: Path) -> dict:
        """Load TXT and optionally split into pages"""
        text = file_path.read_text(encoding="utf-8")
        
        # Split by form feed character (page break) if present
        if '\f' in text or '\x0c' in text:
            page_texts = text.split('\f')
            pages = [
                {"page": i + 1, "text": page_text.strip()}
                for i, page_text in enumerate(page_texts)
                if page_text.strip()
            ]
        elif len(text) > 3000:
            # Split long text files into approximate pages
            pages = DocumentLoader._split_by_length(text, file_path.name)
        else:
            # Short text file - single page
            pages = [{"page": 1, "text": text}]

        return {
            "filename": file_path.name,
            "type": "txt",
            "pages": pages
        }